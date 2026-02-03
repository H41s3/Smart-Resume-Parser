"""Document text extraction service supporting PDF and DOCX."""

import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
from pathlib import Path


class DocumentExtractor:
    """Extract text content from PDF and DOCX files."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

    @classmethod
    def extract_text(cls, file_path: str | Path) -> str:
        """
        Extract text from a document file (PDF or DOCX).

        Args:
            file_path: Path to the document file.

        Returns:
            Extracted text content.

        Raises:
            ValueError: If the file type is not supported or cannot be read.
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return cls._extract_from_pdf_file(path)
        elif ext == ".docx":
            return cls._extract_from_docx_file(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @classmethod
    def extract_text_from_bytes(cls, content: bytes, file_type: str) -> str:
        """
        Extract text from document bytes.

        Args:
            content: File content as bytes.
            file_type: File extension (e.g., '.pdf', '.docx').

        Returns:
            Extracted text content.

        Raises:
            ValueError: If the file type is not supported or content is invalid.
        """
        ext = file_type.lower() if file_type.startswith(".") else f".{file_type.lower()}"

        if ext == ".pdf":
            return cls._extract_from_pdf_bytes(content)
        elif ext == ".docx":
            return cls._extract_from_docx_bytes(content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    def _extract_from_pdf_file(file_path: Path) -> str:
        """Extract text from a PDF file."""
        try:
            doc = fitz.open(file_path)
            text_parts = []

            for page_num in range(doc.page_count):
                page = doc[page_num]
                text_parts.append(page.get_text())

            doc.close()
            return "\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def _extract_from_pdf_bytes(content: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []

            for page_num in range(doc.page_count):
                page = doc[page_num]
                text_parts.append(page.get_text())

            doc.close()
            return "\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def _extract_from_docx_file(file_path: Path) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def _extract_from_docx_bytes(content: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            doc = Document(BytesIO(content))
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
